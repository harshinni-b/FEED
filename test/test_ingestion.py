import json

from docx import Document

from src.ingestion.loader import DocumentParser


def create_fixture(path) -> None:
    document = Document()
    document.core_properties.title = "Test engineering document"
    document.add_heading("Design Basis", level=1)
    document.add_paragraph("The design basis paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tag"
    table.cell(0, 1).text = "Service"
    table.cell(1, 0).text = "P-101"
    table.cell(1, 1).text = "Feed"
    document.save(path)


def test_parse_document_preserves_content(tmp_path) -> None:
    source = tmp_path / "sample.docx"
    create_fixture(source)

    record = DocumentParser().parse_document(source)

    assert record["filename"] == "sample.docx"
    assert record["metadata"]["core_properties"]["title"] == "Test engineering document"
    assert record["headings"] == [{"text": "Design Basis", "level": 1}]
    assert record["paragraphs"][1]["text"] == "The design basis paragraph."
    assert record["tables"][0]["rows"] == [["Tag", "Service"], ["P-101", "Feed"]]


def test_process_all_documents_writes_json_per_docx(tmp_path) -> None:
    raw = tmp_path / "raw"
    extracted = tmp_path / "extracted"
    raw.mkdir()
    create_fixture(raw / "one.docx")
    create_fixture(raw / "two.docx")

    output_paths = DocumentParser().process_all_documents(raw, extracted)

    assert [path.name for path in output_paths] == ["one.json", "two.json"]
    assert json.loads((extracted / "one.json").read_text(encoding="utf-8"))["filename"] == "one.docx"